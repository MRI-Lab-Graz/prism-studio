"""
Render a PRISM survey template as a Word (.docx) paper-pencil questionnaire.

Requires python-docx (optional dependency).
"""
from __future__ import annotations

import io
import random
from datetime import datetime
from typing import Optional


# ── Helpers ──────────────────────────────────────────────────────────

def _get_localized(value, lang: str) -> str:
    if not value:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        if lang in value:
            return str(value[lang])
        for v in value.values():
            if isinstance(v, str) and v.strip():
                return v
    return str(value) if value else ""


def _item_keys(template: dict) -> list[str]:
    skip = {
        "Study", "Technical", "Scoring", "LimeSurvey",
        "MatrixGrouping", "TemplateVersion",
    }
    keys = []
    for k, v in template.items():
        if k in skip or not isinstance(v, dict):
            continue
        if "Description" not in v:
            continue
        keys.append(k)
    keys.sort()
    return keys


def _safe_num(val):
    try:
        return (0, float(val))
    except (ValueError, TypeError):
        return (1, str(val))


def _detect_question_type(item: dict) -> str:
    ls_type = (item.get("LimeSurvey") or {}).get("questionType", "")
    if ls_type == "!":
        return "dropdown"
    if ls_type in ("N", "K"):
        return "numerical"
    if ls_type == "S":
        return "short-text"
    if ls_type in ("T", "U"):
        return "long-text"
    if ls_type == "D":
        return "date"
    if ls_type in ("*", "X"):
        return "hidden"
    input_type = item.get("InputType", "")
    if input_type == "slider":
        return "slider"
    if input_type == "dropdown":
        return "dropdown"
    if input_type == "numerical":
        return "numerical"
    if input_type == "calculated":
        return "hidden"
    if input_type == "text":
        tc = item.get("TextConfig") or {}
        return "long-text" if tc.get("multiline") else "short-text"
    for vs in item.get("VariantScales") or []:
        st = (vs or {}).get("ScaleType", "")
        if st in ("vas", "visual-analogue"):
            return "slider"
    levels = item.get("Levels")
    if levels and isinstance(levels, dict) and len(levels) > 0:
        return "dropdown" if len(levels) > 10 else "radio"
    return "short-text"


def _get_levels(item: dict, variant_id: Optional[str] = None) -> dict:
    if variant_id:
        for vs in item.get("VariantScales") or []:
            if (vs or {}).get("VariantID") == variant_id:
                lvls = vs.get("Levels")
                if lvls and isinstance(lvls, dict) and len(lvls) > 0:
                    return lvls
    return item.get("Levels") or {}


def _levels_signature(levels: dict, lang: str) -> Optional[str]:
    if not levels or not isinstance(levels, dict) or len(levels) < 2:
        return None
    keys = sorted(levels.keys(), key=_safe_num)
    return "|".join(f"{k}={_get_localized(levels[k], lang)}" for k in keys)


# ── Main renderer ────────────────────────────────────────────────────

def render_questionnaire_docx(
    template: dict,
    language: str = "en",
    variant_id: Optional[str] = None,
    options: Optional[dict] = None,
) -> io.BytesIO:
    """
    Render a PRISM survey template as a professional paper-pencil Word document.

    Options dict can include:
        show_participant_id (bool): Add participant ID field at top (default True)
        show_date_field (bool): Add date field at top (default True)
        show_item_codes (bool): Show item codes like GAD701 (default False)
        show_study_info (bool): Show authors/year/citation (default True)
        randomize_items (bool): Randomize item order (default False)
        random_seed (int): Seed for reproducible randomization
        header_repeat_every (int): Repeat scale header every N items in matrix (0=never, default 0)
        font_size (int): Base font size in pt (default 10)
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install it with: pip install python-docx"
        )

    opts = options or {}
    show_pid = opts.get("show_participant_id", True)
    show_date = opts.get("show_date_field", True)
    show_codes = opts.get("show_item_codes", False)
    show_study_info = opts.get("show_study_info", True)
    randomize = opts.get("randomize_items", False)
    random_seed = opts.get("random_seed", None)
    header_repeat = opts.get("header_repeat_every", 0)
    base_font = opts.get("font_size", 10)
    item_col_pct = opts.get("item_column_pct", 55) / 100.0  # default 55%

    doc = Document()

    # -- Page setup
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    study = template.get("Study") or {}
    GRAY = RGBColor(0x66, 0x66, 0x66)
    LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    GREEN = RGBColor(0x1F, 0x8B, 0x5C)
    BG_LIGHT = "F5F5F5"
    BG_WHITE = "FFFFFF"

    def _set_cell_shading(cell, color_hex):
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        })
        shading.append(shd)

    # ── Title block ──────────────────────────────────────────────────

    title_text = _get_localized(study.get("OriginalName"), language)
    short_name = study.get("ShortName", "")

    if title_text:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.space_after = Pt(2)
        run = tp.add_run(title_text)
        run.bold = True
        run.font.size = Pt(base_font + 6)
        run.font.color.rgb = BLACK
        if short_name:
            run2 = tp.add_run(f"\n({short_name})")
            run2.font.size = Pt(base_font + 2)
            run2.font.color.rgb = GRAY

    # Study info (authors, year)
    if show_study_info:
        meta_parts = []
        authors = study.get("Authors")
        if authors:
            meta_parts.append(
                ", ".join(str(a) for a in authors)
                if isinstance(authors, list) else str(authors)
            )
        year = study.get("Year")
        if year:
            meta_parts.append(str(year))
        if meta_parts:
            bp = doc.add_paragraph()
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.space_after = Pt(4)
            br = bp.add_run(" — ".join(meta_parts))
            br.font.size = Pt(base_font - 2)
            br.font.color.rgb = GRAY

    # ── Participant ID / Date fields ─────────────────────────────────

    if show_pid or show_date:
        id_table = doc.add_table(rows=1, cols=2)
        id_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        id_table.autofit = True

        if show_pid:
            cell = id_table.cell(0, 0)
            p = cell.paragraphs[0]
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            r = p.add_run("Participant ID:  ")
            r.bold = True
            r.font.size = Pt(base_font)
            line = p.add_run("_" * 25)
            line.font.size = Pt(base_font)
            line.font.color.rgb = LIGHT_GRAY

        if show_date:
            cell = id_table.cell(0, 1)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            r = p.add_run("Date:  ")
            r.bold = True
            r.font.size = Pt(base_font)
            line = p.add_run("____ / ____ / ________")
            line.font.size = Pt(base_font)
            line.font.color.rgb = LIGHT_GRAY

        # Remove table borders
        for row in id_table.rows:
            for cell in row.cells:
                for border_name in ("top", "bottom", "left", "right"):
                    border = cell._element.get_or_add_tcPr().makeelement(
                        qn(f"w:{border_name}"),
                        {qn("w:val"): "none", qn("w:sz"): "0"},
                    )
                    cell._element.get_or_add_tcPr().append(border)

        doc.add_paragraph().space_before = Pt(2)

    # ── Instructions ─────────────────────────────────────────────────

    instructions = _get_localized(study.get("Instructions"), language)
    if instructions:
        # Bordered instruction box
        instr_table = doc.add_table(rows=1, cols=1)
        instr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = instr_table.cell(0, 0)
        _set_cell_shading(cell, "F0F7F4")
        p = cell.paragraphs[0]
        p.space_before = Pt(6)
        p.space_after = Pt(6)
        r = p.add_run(instructions)
        r.italic = True
        r.font.size = Pt(base_font)
        r.font.color.rgb = BLACK

        # Set cell borders to green left bar
        tc_pr = cell._element.get_or_add_tcPr()
        for side in ("top", "bottom", "right"):
            el = tc_pr.makeelement(
                qn(f"w:{side}"),
                {qn("w:val"): "single", qn("w:sz"): "4", qn("w:color"): "CCCCCC"},
            )
            tc_pr.append(el)
        left_el = tc_pr.makeelement(
            qn("w:left"),
            {qn("w:val"): "single", qn("w:sz"): "24", qn("w:color"): "1F8B5C"},
        )
        tc_pr.append(left_el)

        doc.add_paragraph().space_before = Pt(4)

    # ── Build visible items ──────────────────────────────────────────

    all_keys = _item_keys(template)
    visible = []
    for key in all_keys:
        item = template[key]
        q_type = _detect_question_type(item)
        if q_type == "hidden":
            continue
        applicable = item.get("ApplicableVersions")
        if variant_id and isinstance(applicable, list) and len(applicable) > 0:
            if variant_id not in applicable:
                continue
        levels = _get_levels(item, variant_id)
        visible.append((key, item, q_type, levels))

    if randomize:
        rng = random.Random(random_seed)
        rng.shuffle(visible)

    # ── Group consecutive radio items with same scale into matrix ────

    groups = []
    i = 0
    while i < len(visible):
        key, item, q_type, levels = visible[i]
        if q_type == "radio" and levels:
            sig = _levels_signature(levels, language)
            if sig:
                block = [(key, item, levels)]
                j = i + 1
                while j < len(visible):
                    nk, ni, nt, nl = visible[j]
                    if nt == "radio" and _levels_signature(nl, language) == sig:
                        block.append((nk, ni, nl))
                        j += 1
                    else:
                        break
                if len(block) >= 2:
                    groups.append(("matrix", block))
                    i = j
                    continue
        groups.append(("single", [(key, item, levels)]))
        i += 1

    # ── Render groups ────────────────────────────────────────────────

    question_num = 0

    for group_type, block in groups:

        if group_type == "matrix":
            levels = block[0][2]
            sorted_lk = sorted(levels.keys(), key=_safe_num)
            n_cols = len(sorted_lk)

            # Create table: cols = item_text + one col per level
            table = doc.add_table(rows=0, cols=1 + n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # Set column widths using configurable item column percentage
            total_width_cm = 17.0
            item_col_width = Cm(total_width_cm * item_col_pct)
            level_col_width = Cm(total_width_cm * (1 - item_col_pct) / max(n_cols, 1))

            # Apply widths via table grid
            for col_idx, col in enumerate(table.columns):
                col.width = item_col_width if col_idx == 0 else level_col_width

            def _apply_row_widths(row_obj):
                row_obj.cells[0].width = item_col_width
                for ci in range(1, len(row_obj.cells)):
                    row_obj.cells[ci].width = level_col_width

            def _add_header_row(tbl):
                hr = tbl.add_row()
                _apply_row_widths(hr)
                # Item column header
                hc = hr.cells[0]
                _set_cell_shading(hc, "E8E8E8")
                hp = hc.paragraphs[0]
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                hp.space_before = Pt(3)
                hp.space_after = Pt(3)
                # Level headers
                for ci, lk in enumerate(sorted_lk):
                    hc = hr.cells[ci + 1]
                    _set_cell_shading(hc, "E8E8E8")
                    hp = hc.paragraphs[0]
                    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    hp.space_before = Pt(3)
                    hp.space_after = Pt(3)
                    lv_text = _get_localized(levels[lk], language)
                    r = hp.add_run(lv_text)
                    r.bold = True
                    r.font.size = Pt(base_font - 2)
                return hr

            # First header
            _add_header_row(table)

            for idx, (key, item, _lvls) in enumerate(block):
                question_num += 1

                # Repeat header every N items
                if header_repeat > 0 and idx > 0 and idx % header_repeat == 0:
                    _add_header_row(table)

                row = table.add_row()
                _apply_row_widths(row)
                # Alternate row shading
                bg = BG_LIGHT if idx % 2 == 0 else BG_WHITE

                # Item cell
                item_cell = row.cells[0]
                _set_cell_shading(item_cell, bg)
                p = item_cell.paragraphs[0]
                p.space_before = Pt(2)
                p.space_after = Pt(2)

                nr = p.add_run(f"{question_num}. ")
                nr.bold = True
                nr.font.size = Pt(base_font)

                if show_codes:
                    cr = p.add_run(f"[{key}] ")
                    cr.font.size = Pt(base_font - 3)
                    cr.font.color.rgb = LIGHT_GRAY

                desc = _get_localized(item.get("Description"), language)
                dr = p.add_run(desc)
                dr.font.size = Pt(base_font)

                if item.get("Reversed"):
                    rr = p.add_run(" (R)")
                    rr.font.size = Pt(base_font - 3)
                    rr.font.color.rgb = LIGHT_GRAY
                    rr.italic = True

                # Radio cells with circles
                for ci in range(n_cols):
                    rc = row.cells[ci + 1]
                    _set_cell_shading(rc, bg)
                    rp = rc.paragraphs[0]
                    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rp.space_before = Pt(2)
                    rp.space_after = Pt(2)
                    circle = rp.add_run("○")
                    circle.font.size = Pt(base_font + 2)

            # Style table borders
            _style_matrix_table(table)
            doc.add_paragraph().space_before = Pt(6)

        else:
            # Single item
            key, item, levels = block[0]
            q_type = _detect_question_type(item)
            question_num += 1

            p = doc.add_paragraph()
            p.space_before = Pt(8)
            p.space_after = Pt(2)

            nr = p.add_run(f"{question_num}. ")
            nr.bold = True
            nr.font.size = Pt(base_font)

            if show_codes:
                cr = p.add_run(f"[{key}] ")
                cr.font.size = Pt(base_font - 3)
                cr.font.color.rgb = LIGHT_GRAY

            desc = _get_localized(item.get("Description"), language)
            dr = p.add_run(desc)
            dr.font.size = Pt(base_font)

            if item.get("Reversed"):
                rr = p.add_run(" (R)")
                rr.font.size = Pt(base_font - 3)
                rr.font.color.rgb = LIGHT_GRAY
                rr.italic = True

            # Item instructions
            item_instr = _get_localized(item.get("Instructions"), language)
            if item_instr:
                ip = doc.add_paragraph()
                ip.space_before = Pt(0)
                ip.space_after = Pt(2)
                ip.paragraph_format.left_indent = Cm(0.5)
                ir = ip.add_run(item_instr)
                ir.font.size = Pt(base_font - 1)
                ir.italic = True
                ir.font.color.rgb = GRAY

            # Render response based on type
            if q_type == "radio" and levels:
                # Single radio item (not in matrix) - inline circles
                sorted_lk = sorted(levels.keys(), key=_safe_num)
                rp = doc.add_paragraph()
                rp.paragraph_format.left_indent = Cm(0.5)
                rp.space_before = Pt(2)
                rp.space_after = Pt(4)
                for li, lk in enumerate(sorted_lk):
                    lv_text = _get_localized(levels[lk], language)
                    if li > 0:
                        sep = rp.add_run("     ")
                        sep.font.size = Pt(base_font)
                    circle = rp.add_run("○ ")
                    circle.font.size = Pt(base_font + 1)
                    lr = rp.add_run(lv_text)
                    lr.font.size = Pt(base_font - 1)

            elif q_type == "dropdown" and levels:
                sorted_lk = sorted(levels.keys(), key=_safe_num)
                for lk in sorted_lk:
                    lv_text = _get_localized(levels[lk], language)
                    rp = doc.add_paragraph()
                    rp.paragraph_format.left_indent = Cm(0.8)
                    rp.space_before = Pt(0)
                    rp.space_after = Pt(0)
                    rp.add_run("□ ").font.size = Pt(base_font)
                    lr = rp.add_run(lv_text)
                    lr.font.size = Pt(base_font - 1)

            elif q_type == "numerical":
                np_p = doc.add_paragraph()
                np_p.paragraph_format.left_indent = Cm(0.5)
                np_p.space_before = Pt(2)
                np_p.add_run("___________").font.size = Pt(base_font)
                parts = []
                if item.get("MinValue") is not None:
                    parts.append(f"Min: {item['MinValue']}")
                if item.get("MaxValue") is not None:
                    parts.append(f"Max: {item['MaxValue']}")
                if item.get("Unit"):
                    parts.append(item["Unit"])
                if parts:
                    hint = np_p.add_run(f"  ({', '.join(parts)})")
                    hint.font.size = Pt(base_font - 2)
                    hint.font.color.rgb = GRAY

            elif q_type == "slider":
                sc = item.get("SliderConfig") or {}
                min_v = sc.get("min", item.get("MinValue", 0))
                max_v = sc.get("max", item.get("MaxValue", 100))
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Cm(0.5)
                sp.space_before = Pt(4)
                sr = sp.add_run(f"{min_v}  |")
                sr.font.size = Pt(base_font - 1)
                bar = sp.add_run("━" * 35)
                bar.font.size = Pt(base_font - 1)
                bar.font.color.rgb = LIGHT_GRAY
                sr2 = sp.add_run(f"|  {max_v}")
                sr2.font.size = Pt(base_font - 1)

            elif q_type == "long-text":
                for _ in range(3):
                    lp = doc.add_paragraph()
                    lp.paragraph_format.left_indent = Cm(0.5)
                    lp.space_before = Pt(0)
                    lp.space_after = Pt(0)
                    lr = lp.add_run("_" * 75)
                    lr.font.size = Pt(base_font - 1)
                    lr.font.color.rgb = LIGHT_GRAY

            elif q_type == "date":
                dp = doc.add_paragraph()
                dp.paragraph_format.left_indent = Cm(0.5)
                dp.space_before = Pt(2)
                dp.add_run("____ / ____ / ________").font.size = Pt(base_font)

            else:  # short-text
                tp_p = doc.add_paragraph()
                tp_p.paragraph_format.left_indent = Cm(0.5)
                tp_p.space_before = Pt(2)
                lr = tp_p.add_run("_" * 55)
                lr.font.size = Pt(base_font)
                lr.font.color.rgb = LIGHT_GRAY

    # ── Footer ───────────────────────────────────────────────────────

    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"Generated by PRISM Studio — {datetime.now().strftime('%Y-%m-%d')} — "
        f"{question_num} items"
    )
    fr.font.size = Pt(7)
    fr.font.color.rgb = LIGHT_GRAY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _style_matrix_table(table):
    """Apply clean borders to a matrix table."""
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{side}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "CCCCCC",
        })
        borders.append(el)
    tbl_pr.append(borders)
