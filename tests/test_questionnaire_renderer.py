"""Tests for questionnaire_renderer module (Word export)."""

import io
import json
import pytest

pytest.importorskip("docx")

# Ensure app/ is importable
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "app"))


@pytest.fixture
def likert_template():
    return {
        "Study": {
            "OriginalName": {"en": "Test Questionnaire", "de": "Test-Fragebogen"},
            "ShortName": "TQ",
            "Authors": ["Test Author"],
            "Year": 2026,
            "Instructions": {
                "en": "Please answer all questions.",
                "de": "Bitte beantworten Sie alle Fragen.",
            },
        },
        "Technical": {"StimulusType": "Questionnaire", "Language": ""},
        "TQ01": {
            "Description": {"en": "I feel happy", "de": "Ich bin gluecklich"},
            "Levels": {
                "0": {"en": "Never"},
                "1": {"en": "Sometimes"},
                "2": {"en": "Always"},
            },
            "DataType": "integer",
        },
        "TQ02": {
            "Description": {"en": "I feel calm", "de": "Ich bin ruhig"},
            "Levels": {
                "0": {"en": "Never"},
                "1": {"en": "Sometimes"},
                "2": {"en": "Always"},
            },
            "DataType": "integer",
            "Reversed": True,
        },
        "TQ03": {
            "Description": {"en": "I feel strong", "de": "Ich bin stark"},
            "Levels": {
                "0": {"en": "Never"},
                "1": {"en": "Sometimes"},
                "2": {"en": "Always"},
            },
            "DataType": "integer",
        },
    }


@pytest.fixture
def mixed_template():
    return {
        "Study": {
            "OriginalName": "Mixed Types",
            "ShortName": "MIX",
        },
        "Technical": {"StimulusType": "Questionnaire"},
        "M01": {
            "Description": "How old are you?",
            "InputType": "numerical",
            "MinValue": 18,
            "MaxValue": 99,
        },
        "M02": {
            "Description": "Describe your feelings",
            "InputType": "text",
            "TextConfig": {"multiline": True, "rows": 3},
        },
        "M03": {
            "Description": "Rate your pain",
            "InputType": "slider",
            "SliderConfig": {"min": 0, "max": 100},
        },
        "M04": {
            "Description": "Select your country",
            "InputType": "dropdown",
            "Levels": {str(i): f"Country {i}" for i in range(15)},
        },
    }


def test_render_likert_produces_valid_docx(likert_template):
    from src.questionnaire_renderer import render_questionnaire_docx

    buf = render_questionnaire_docx(likert_template, language="en")
    assert isinstance(buf, io.BytesIO)
    assert buf.getvalue()[:4] == b"PK\x03\x04"  # ZIP magic (docx is a ZIP)

    from docx import Document

    doc = Document(buf)
    # Should have tables: PID/Date, Instructions, Matrix
    assert len(doc.tables) >= 2
    # Matrix table should have header + 3 items
    matrix = [t for t in doc.tables if len(t.columns) > 2]
    assert len(matrix) == 1
    assert len(matrix[0].rows) == 4  # 1 header + 3 items


def test_render_likert_german(likert_template):
    from src.questionnaire_renderer import render_questionnaire_docx
    from docx import Document

    buf = render_questionnaire_docx(likert_template, language="de")
    doc = Document(buf)
    # Title should be German
    title_text = doc.paragraphs[0].text
    assert "Test-Fragebogen" in title_text


def test_render_with_options(likert_template):
    from src.questionnaire_renderer import render_questionnaire_docx
    from docx import Document

    buf = render_questionnaire_docx(
        likert_template,
        language="en",
        options={
            "show_participant_id": False,
            "show_date_field": False,
            "show_study_info": False,
            "show_item_codes": True,
            "font_size": 12,
            "item_column_pct": 60,
        },
    )
    doc = Document(buf)
    # No PID/Date table (2 cols), only instructions + matrix
    pid_tables = [t for t in doc.tables if len(t.columns) == 2]
    assert len(pid_tables) == 0


def test_render_mixed_types(mixed_template):
    from src.questionnaire_renderer import render_questionnaire_docx
    from docx import Document

    buf = render_questionnaire_docx(mixed_template, language="en")
    doc = Document(buf)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "How old are you?" in full_text
    assert "Min: 18" in full_text
    assert "Describe your feelings" in full_text
    assert "Rate your pain" in full_text


def test_render_randomized(likert_template):
    from src.questionnaire_renderer import render_questionnaire_docx
    from docx import Document

    buf1 = render_questionnaire_docx(
        likert_template,
        language="en",
        options={"randomize_items": True, "random_seed": 42},
    )
    buf2 = render_questionnaire_docx(
        likert_template,
        language="en",
        options={"randomize_items": True, "random_seed": 42},
    )
    # Same seed = same output
    assert buf1.getvalue() == buf2.getvalue()

    buf3 = render_questionnaire_docx(
        likert_template,
        language="en",
        options={"randomize_items": True, "random_seed": 99},
    )
    # Different seed = different output (items reordered)
    # Note: with only 3 items there's a chance of same order, but very unlikely
    doc1 = Document(io.BytesIO(buf1.getvalue()))
    doc3 = Document(io.BytesIO(buf3.getvalue()))
    # At least verify both produce valid docs
    assert len(doc1.tables) > 0
    assert len(doc3.tables) > 0


def test_hidden_items_excluded(likert_template):
    from src.questionnaire_renderer import render_questionnaire_docx
    from docx import Document

    likert_template["HIDDEN01"] = {
        "Description": "Should not appear",
        "InputType": "calculated",
        "Calculation": {"formula": "TQ01 + TQ02"},
    }
    buf = render_questionnaire_docx(likert_template, language="en")
    doc = Document(buf)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Should not appear" not in full_text
